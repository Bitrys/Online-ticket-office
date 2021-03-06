"""
This program is made to order by Жмых Airlines, and is provided as a project on Yandex Lyceum by Qt.
This program is able to book a seat on the flight, print the ticket and save the order in the database.
REMEMBER! No matter how, no matter where, the main thing is together! Жмых Airlines!

Error codes:
E001 — error in the func "start_bought_process"
E002 — error in the func "print_ticket"
E003 — error in the func "add_ticket_at_database"
E004 —  unknown error (fatal)
"""

import sqlite3
import sys
import random

import barcode
import cython
import simpleaudio as sa
from PyQt5 import QtCore, QtGui
from PyQt5.QtWidgets import QApplication, QWidget
from PyQt5 import uic
from docx import Document
from docx.shared import Inches
from barcode.writer import ImageWriter


class Main(QWidget):
    def __init__(self):
        """
        Initialized program.
        """

        super(Main, self).__init__()
        uic.loadUi('mainUI.ui', self)

        self.choise_class.setEditable(True)
        self.choise_class.lineEdit().setAlignment(QtCore.Qt.AlignCenter)
        self.choise_class.lineEdit().setFont(QtGui.QFont('Roboto Slab', 14))
        self.bought.clicked.connect(self.start_bought_process)
        self.choise_class.lineEdit().setDisabled(True)

        wave_obj = sa.WaveObject.from_wave_file('data/theme_app.wav')
        wave_obj.play()

    def start_bought_process(self):
        """
        Check and progress user input and buy a ticket.
        """

        self.total_sum_output.setText('')

        id_ticket = random.randint(123456789012, 9999999999999)
        class_flight = self.choise_class.currentText()
        address_from, address_to, time = self.froms.text(), self.to.text(), self.when.text()
        fio, passport_data = self.fio.text(), self.passport.text()
        promocode = self.promos.text().lower()

        if address_from == '' or address_to == '' or time == '' or fio == '' or passport_data == '':
            self.status.setStyleSheet('background: red;')
            self.status.setText('Fill all lines!')
            self.total_sum_output.setText('')
        else:
            self.status.setStyleSheet('background: green;')
            self.status.setText('Ok!')
            try:
                self.count_sum_bought(class_flight, promocode)
                self.print_ticket(id_ticket, fio, address_from, address_to, time)
                self.add_ticket_at_database(id_ticket, address_from, address_to, time, fio, passport_data, class_flight)
            except Exception as e:
                if str(e) == 'Error code: E001!' or str(e) == 'Error code: E002!' or str(e) == 'Error code: E003!':
                    self.status.setStyleSheet('background: red;')
                    self.status.setText(str(e))
                    self.total_sum_output.setText('')
                else:
                    self.status.setStyleSheet('background: red;')
                    self.status.setText('Fatal error code: E004!')
                    self.total_sum_output.setText('')
                    print(str(e))
                    print('Please, report this error, as well as the latest actions, to the system administrator!')

    def count_sum_bought(self, class_flight, promocode):
        """
        Counted sum of bought.
        :param class_flight: the type of ticket
        :param promocode: promo code company
        """

        try:
            coefficient = {'Первый класс': 2.5, 'VIP': 2, 'Бизнес-класс': 1.5, 'Стандарт': 1}
            price = random.randint(1000, 20000) * coefficient[class_flight]

            discount = {
                'free plz': lambda _: 0,
                'скидка': lambda _: _ / 2,
                'discount': lambda _: _ / 2,
                'zmih air': lambda _: _ / 1.5,
                'жмых эйр': lambda _: _ / 1.5,
                'не важно где, не важно как, главное вместе': lambda _: _ / 1.8,
                'яндекс лицей': lambda _: _ / 8,
                "123', '123'); drop table main; --": lambda _: _ * 8
            }

            if promocode in discount:
                self.total_sum_output.setText(f'{discount[promocode](price)} ₽')
            else:
                self.total_sum_output.setText(f'{price} ₽')
        except Exception as e:
            print('Message to the system administrator:')
            print(e)
            raise Exception('Error code: E001!')

    @staticmethod
    def print_ticket(id_ticket, fio, address_from, address_to, time):
        """
        Printed ticket at ticket.docx
        :param time: flight date
        :param id_ticket: unique ticket id
        :param fio: the initials of the passenger
        :param address_from: airport of departure
        :param address_to: arrival airport
        """

        try:
            doc = Document()
            doc.add_picture('data/zmih_logo.png', width=Inches(4))
            doc.add_heading(f'Жмых эйр: билет № {random.randint(10, 999999)}', 0)

            doc.add_heading('Пассажир:', level=1)
            doc.add_paragraph(fio)

            doc.add_heading('Рейс №:', level=1)
            doc.add_paragraph(f'ЖМА — {random.randint(100, 999)}')

            doc.add_heading('Уникальный ID билета:', level=1)
            doc.add_paragraph(str(id_ticket))

            doc.add_heading('Аэропорт вылета:', level=1)
            doc.add_paragraph(f'[{address_from[:3].upper()}] — {address_from}')

            doc.add_heading('Аэропорт прилёта:', level=1)
            doc.add_paragraph(f'[{address_to[:3].upper()}] — {address_to}')

            doc.add_heading('Дата вылета:', level=1)
            doc.add_paragraph(time)

            ean = barcode.codex.Code39(str(id_ticket), writer=ImageWriter(), add_checksum=False)  # created barcode
            ean.save(f'data/{id_ticket}')  # save barcode
            doc.add_picture(f'data/{id_ticket}.png', width=Inches(2))

            doc.save('ticket.docx')
        except Exception as e:
            print('Message to the system administrator:')
            print(e)
            raise Exception('Error code: E002!')

    @staticmethod
    def add_ticket_at_database(id_ticket, address_from, address_to, time, fio, passport_data, class_flight):
        """
        Added ticket at database.
        :param id_ticket: unique ticket id
        :param address_from: airport of departure
        :param address_to: arrival airport
        :param time: flight date
        :param fio: the initials of the passenger
        :param passport_data: passport and registration data
        :param class_flight: the type of ticket
        """

        try:
            base = sqlite3.connect('data/base_of_tickets.db')
            cursor = base.cursor()
            req = f'''INSERT INTO main(id,fio,from_airport,to_airport,passport_data,time_of_flight,type_of_seat, \
                    barcode_path) VALUES (?, ?, ?, ?, ?, ?, ?, ?) '''
            cursor.execute(req, (id_ticket, fio, address_from, address_to, passport_data, time, class_flight,
                                 f'data/{id_ticket}.png'))
            base.commit()
        except Exception as e:
            print('Message to the system administrator:')
            print(e)
            raise Exception('Error code: E003!')


def except_hook(cls, exception, traceback):
    sys.__excepthook__(cls, exception, traceback)


if __name__ == "__main__":
    # sys args of program

    app = QApplication(sys.argv)
    ex = Main()
    ex.show()
    sys.excepthook = except_hook
    sys.exit(app.exec())
